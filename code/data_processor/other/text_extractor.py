import docx
from docx import Document
import xlsxwriter as xl
import pandas as pd
import spacy
import textract
nlp = spacy.load('en_core_web_sm')


keyword_dic = ['防火墙', '耐火等级', '2h', '12次/h']
file_list = ['805_硬泡聚氨酯保温防水工程技术规范GB_50404_2017.docx',
            '806-外墙内保温工程技术规程JGJT 261-2011 (1).docx',
            '建筑设计防火规范2018.doc']

def read_to_excel(filedir, file_name, keywords_list):
    try:
        f = Document(filedir + file_name)
        paras_num = len(f.paragraphs)
        print('The current file paragraph number is %d' % paras_num)
        para_list = []
        for para in f.paragraphs:
            para_list.append(para.text)

        release_time = ''
        for word in para_list:
            # print(word)
            if '年' in word and '月' in word and '日' in word:
                release_time = word
        if not release_time:
            release_time = '没有发布时间'


        f_text = ''.join(para_list)



        keyword_count = {}
        count = 0
        for keyword in keywords_list:
            for words in f_text.split():
                if keyword in words:
                    count += 1
            keyword_count[keyword] = count

        sent_list = [str(i) for i in list(nlp(f_text).sents)]

        content_list = []
        content_list.append(file_name)
        content_list.append(release_time)
        content_list.append('关键词出现次数统计为：' + str(keyword_count))



        for keyword in keywords_list:
            keyword_dic = {}
            for idx, sentence in enumerate(sent_list):
                if keyword in sentence:
                    if keyword not in keyword_dic:
                        keyword_dic[keyword] = [''.join(sent_list[idx-1:idx+2])]
                    else:
                        keyword_dic[keyword].append(''.join(sent_list[idx-1:idx+2]))
            content_list.append(keyword_dic)

        print(content_list)
        print(len(content_list))

        # with xl.Workbook(filedir + file_name + '.xlsx') as workbook:
        #     worksheet = workbook.add_worksheet()
        #
        #     for col_num, data in enumerate(content_list):
        #         worksheet.write_row(col_num, 0, str(data))

        df = pd.DataFrame(content_list)
        writer = pd.ExcelWriter(filedir + file_name + '.xlsx', engine='xlsxwriter')
        df.to_excel(writer,index=False)
        writer.save()
    except:
        f = textract.process(filedir + file_name)
        print(type(f))
        text = f.decode("utf-8")
        print(type(text))
        para_list = text.split('\n')
        para_list = [i.strip() for i in para_list if len(i)!= 0]
        para_list = [i.strip() for i in para_list]
        print('current file length is %d' % len(para_list))
        release_time = ''
        for word in para_list:
            if '年' in word and '月' in word and '日' in word:
                release_time = word
        if not release_time:
            release_time = '没有发布时间'

        f_text = ''.join(para_list)

        keyword_count = {}
        count = 0
        for keyword in keywords_list:
            for words in f_text.split():
                if keyword in words:
                    count += 1
            keyword_count[keyword] = count

        sent_list = [str(i) for i in list(nlp(f_text).sents)]

        content_list = []
        content_list.append('文件名称：' + file_name)
        content_list.append('发布时间为：' + release_time)
        content_list.append('关键词出现次数统计为：' + str(keyword_count))



        for keyword in keywords_list:
            keyword_dic = {}
            for idx, sentence in enumerate(sent_list):
                if keyword in sentence:
                    if keyword not in keyword_dic:
                        keyword_dic[keyword] = [''.join(sent_list[idx-1:idx+2])]
                    else:
                        keyword_dic[keyword].append(''.join(sent_list[idx-1:idx+2]))
            content_list.append(keyword_dic)

        # print(len(content_list))


        # with xl.Workbook(filedir + file_name + '.xlsx') as workbook:
        #     worksheet = workbook.add_worksheet()
        #
        #     for col_num, data in enumerate(content_list):
        #         worksheet.write_row(col_num, 0,  str(data))

        df = pd.DataFrame(content_list)
        writer = pd.ExcelWriter(filedir + file_name + '.xlsx', engine='xlsxwriter')
        df.to_excel(writer,index=False)
        writer.save()



if __name__ == '__main__':

    # 把下面文件的路颈换为本地路径即可
    for file in file_list:
        read_to_excel('/home/jade/Documents/', file, keyword_dic)
